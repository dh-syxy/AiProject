"""
Disney RAG helpers: DOCX→Markdown export, image OCR, CLIP image+OCR-text FAISS,
and FAISS on DOCX Markdown via DashScope ``text-embedding-v4``.

CLIP 下载默认使用 ``HF_ENDPOINT=https://hf-mirror.com``（可在环境中覆盖）。
"""

from __future__ import annotations

import mimetypes
import os
import pickle
import re
import shutil
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterator, List, Literal, Optional, Tuple, Union

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

try:
    from PIL import Image, ImageEnhance, ImageFilter, ImageOps
except ImportError:  # pragma: no cover
    Image = None  # type: ignore[misc, assignment]
    ImageEnhance = ImageFilter = ImageOps = None  # type: ignore[misc, assignment]

try:
    import pytesseract  # type: ignore[import-untyped]
except ImportError:  # pragma: no cover
    pytesseract = None  # type: ignore[misc, assignment]

ChunkType = Literal["text", "table"]

# macOS 上 faiss 与其它 OpenMP 库并存时可能触发 OMP Error #15
os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")

# Hugging Face Hub：默认走国内镜像（下载 CLIP 等）；如需官方站，运行前 ``unset HF_ENDPOINT`` 或自行 export 覆盖
os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")

# HTTPS：减轻 macOS/pyenv 下 ``SSL: CERTIFICATE_VERIFY_FAILED``（与 ``pufa-bank-rag`` 等一致）
try:
    import certifi

    os.environ.setdefault("REQUESTS_CA_BUNDLE", certifi.where())
    os.environ.setdefault("SSL_CERT_FILE", certifi.where())
    os.environ.setdefault("CURL_CA_BUNDLE", certifi.where())
except ImportError:
    pass

try:
    import truststore  # type: ignore[import-not-found]

    truststore.inject_into_ssl()
except ImportError:
    pass

# Subdirectory under ``disney_knowledge_base`` for FAISS persistence (DOCX-derived MD only).
_DEFAULT_DISNEY_FAISS_DIRNAME = "disney_docx_faiss"

# Lowercase suffixes (with dot) treated as raster images for OCR.
_IMAGE_SUFFIXES = frozenset(
    {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tif", ".tiff", ".webp"}
)


def _resolve_tesseract_cmd(explicit: Optional[str] = None) -> None:
    """Point pytesseract at the ``tesseract`` binary (PATH or common Homebrew paths)."""
    if explicit:
        pytesseract.pytesseract.tesseract_cmd = explicit  # type: ignore[union-attr]
        return
    if shutil.which("tesseract"):
        return
    for candidate in (
        "/opt/homebrew/bin/tesseract",
        "/usr/local/bin/tesseract",
    ):
        if Path(candidate).is_file():
            pytesseract.pytesseract.tesseract_cmd = candidate  # type: ignore[union-attr]
            return


def _preprocess_image_for_ocr(
    im: "Image.Image",
    *,
    min_long_edge: int,
    max_long_edge: Optional[int],
    contrast: float,
    sharpness: float,
    apply_median: bool,
    binarize: bool,
    threshold: int,
    unsharp_mask: bool = True,
) -> "Image.Image":
    """
    Light cleanup for Tesseract: grayscale, optional upscale, autocontrast, sharpen.
    Optional UnsharpMask helps thin Latin letters next to CJK after upscaling.
    Pillow-only (no OpenCV required).
    """
    if im.mode in ("RGBA", "P"):
        im = im.convert("RGB")
    gray = im.convert("L")
    w, h = gray.size
    long_edge = max(w, h)
    if long_edge < min_long_edge:
        scale = min_long_edge / float(long_edge)
        nw, nh = max(1, int(w * scale)), max(1, int(h * scale))
        _resample = getattr(Image, "Resampling", Image).LANCZOS  # Pillow 9+
        gray = gray.resize((nw, nh), _resample)
    elif max_long_edge is not None and long_edge > max_long_edge:
        scale = max_long_edge / float(long_edge)
        nw, nh = max(1, int(w * scale)), max(1, int(h * scale))
        _resample = getattr(Image, "Resampling", Image).LANCZOS
        gray = gray.resize((nw, nh), _resample)

    assert ImageOps is not None
    gray = ImageOps.autocontrast(gray, cutoff=2)
    gray = ImageEnhance.Contrast(gray).enhance(contrast)  # type: ignore[union-attr]
    gray = ImageEnhance.Sharpness(gray).enhance(sharpness)  # type: ignore[union-attr]
    if unsharp_mask:
        try:
            um = getattr(ImageFilter, "UnsharpMask", None)
            if um is not None:
                gray = gray.filter(um(radius=1.1, percent=125, threshold=2))  # type: ignore[union-attr]
        except Exception:
            pass
    if apply_median:
        gray = gray.filter(ImageFilter.MedianFilter(size=3))  # type: ignore[union-attr]
    if binarize:
        gray = gray.point(lambda p: 255 if p > threshold else 0, mode="1").convert("L")
    return gray


def _build_tesseract_config(
    *,
    oem: int,
    psm: int,
    extra: str,
    config_vars: Optional[Dict[str, Union[str, int]]] = None,
) -> str:
    parts = [f"--oem {int(oem)}", f"--psm {int(psm)}"]
    if config_vars:
        for key, val in config_vars.items():
            parts.append(f"-c {key}={val}")
    if extra.strip():
        parts.append(extra.strip())
    return " ".join(parts)


def _default_zh_en_tesseract_vars(
    *,
    preserve_interword_spaces: bool,
) -> Dict[str, Union[str, int]]:
    """
    Extra Tesseract variables for Chinese + English in one pass (``chi_sim+eng``).
    ``preserve_interword_spaces`` helps English word boundaries in mixed lines.
    """
    out: Dict[str, Union[str, int]] = {}
    if preserve_interword_spaces:
        out["preserve_interword_spaces"] = "1"
    return out


@dataclass
class ChunkRecord:
    """One exported chunk (paragraph-level text segment or one table)."""

    chunk_index: int
    chunk_type: ChunkType
    page: int
    page_source: Literal["break", "estimated", "default"]
    content_preview: str


@dataclass
class DocxExportResult:
    """Per-file conversion summary."""

    docx_filename: str
    docx_path: str
    md_filename: str
    md_path: str
    chunks: List[ChunkRecord] = field(default_factory=list)
    page_note: str = ""


@dataclass
class ImageOcrResult:
    """Per-image OCR export summary."""

    image_filename: str
    image_path: str
    image_type: str
    md_filename: str
    md_path: str
    ocr_text: str


@dataclass
class ClipFaissBuildResult:
    """Summary after building CLIP + FAISS from knowledge-base images."""

    save_dir: str
    clip_model_name: str
    embedding_dim: int
    num_vectors: int
    num_image_vectors: int
    num_text_vectors: int


@dataclass
class RagAnswerResult:
    """Answer from retrieval + DashScope chat model with citation hints."""

    answer: str
    intent: Literal["text_knowledge", "image_knowledge"]
    llm_model: str
    top_contexts: List[Dict[str, Any]]
    prompt_sources_summary: str


@dataclass
class FaissBuildResult:
    """Summary after building FAISS from DOCX-exported Markdown files."""

    save_dir: str
    embedding_model: str
    md_files_used: int
    num_chunks: int
    faiss_store: Any


def _recursive_text_splitter():
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        return RecursiveCharacterTextSplitter
    except ImportError:
        from langchain.text_splitter import RecursiveCharacterTextSplitter  # type: ignore[no-redef]

        return RecursiveCharacterTextSplitter


def build_faiss_from_docx_markdown(
    knowledge_dir: Optional[Union[str, Path]] = None,
    *,
    save_dir: Optional[Union[str, Path]] = None,
    embedding_model: str = "text-embedding-v4",
    chunk_size: int = 800,
    chunk_overlap: int = 150,
    dashscope_api_key: Optional[str] = None,
    verbose: bool = True,
) -> FaissBuildResult:
    """
    Read **top-level** ``*.md`` files under ``disney_knowledge_base`` (same outputs as
    :func:`export_docx_folder_to_markdown`), split into chunks, embed with Alibaba Cloud
    DashScope ``text-embedding-v4``, and persist a LangChain ``FAISS`` index.

    Only Markdown files directly in ``knowledge_dir`` are used (not ``images/*.md`` from OCR).

    Dependencies::

        pip install langchain-community faiss-cpu
        # plus langchain / langchain-text-splitters as required by your LangChain version

    Set environment variable ``DASHSCOPE_API_KEY`` (百炼 / DashScope API Key).

    Args:
        knowledge_dir: Folder containing ``*.md`` from DOCX export. Defaults to
            ``./disney_knowledge_base`` next to this module.
        save_dir: Directory for ``index.faiss``, ``index.pkl``. Defaults to
            ``<knowledge_dir>/disney_docx_faiss``.
        embedding_model: DashScope embedding model name (default ``text-embedding-v4``).
        chunk_size / chunk_overlap: Recursive character splitter settings (Chinese-friendly separators).
        dashscope_api_key: Optional; defaults to ``DASHSCOPE_API_KEY`` env.
        verbose: Print counts and save path.

    Returns:
        :class:`FaissBuildResult` with the in-memory ``FAISS`` store and stats.
    """
    from langchain_community.embeddings import DashScopeEmbeddings
    from langchain_community.vectorstores import FAISS

    api_key = dashscope_api_key or os.getenv("DASHSCOPE_API_KEY")
    if not api_key:
        raise ValueError(
            "请设置环境变量 DASHSCOPE_API_KEY，或通过 dashscope_api_key 传入 DashScope API Key。"
        )

    base = Path(__file__).resolve().parent
    folder = Path(knowledge_dir) if knowledge_dir is not None else base / "disney_knowledge_base"
    folder = folder.resolve()

    if not folder.is_dir():
        raise FileNotFoundError(f"知识库目录不存在: {folder}")

    md_paths = sorted(folder.glob("*.md"))
    if not md_paths:
        raise ValueError(f"目录下没有顶层 .md 文件（请先运行 export_docx_folder_to_markdown）: {folder}")

    RecursiveCharacterTextSplitter = _recursive_text_splitter()
    separators = ["\n\n", "\n", "。", "！", "？", "；", ".", " ", ""]
    splitter = RecursiveCharacterTextSplitter(
        separators=separators,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )

    all_chunks: List[str] = []
    all_metadatas: List[Dict[str, Any]] = []

    for md_path in md_paths:
        raw = md_path.read_text(encoding="utf-8").strip()
        if not raw:
            continue
        pieces = splitter.split_text(raw)
        docx_sidecar = md_path.with_suffix(".docx")
        docx_path_str = str(docx_sidecar.resolve()) if docx_sidecar.is_file() else ""
        for i, piece in enumerate(pieces):
            if not piece.strip():
                continue
            nums, dom = _page_numbers_from_md_export_chunk(piece)
            lo = min(nums) if nums else None
            hi = max(nums) if nums else None
            all_chunks.append(piece)
            all_metadatas.append(
                {
                    "source": md_path.name,
                    "md_path": str(md_path.resolve()),
                    "docx_path": docx_path_str,
                    "chunk_in_file": i,
                    "page": dom,
                    "page_min": lo,
                    "page_max": hi,
                    "page_label": _page_label_from_numbers(nums) or None,
                }
            )

    if not all_chunks:
        raise ValueError("切分后没有任何文本块；请检查 .md 是否为空。")

    embeddings = DashScopeEmbeddings(
        model=embedding_model,
        dashscope_api_key=api_key,
    )
    store = FAISS.from_texts(all_chunks, embeddings, metadatas=all_metadatas)

    out_dir = (
        Path(save_dir).expanduser().resolve()
        if save_dir is not None
        else folder / _DEFAULT_DISNEY_FAISS_DIRNAME
    )
    out_dir.mkdir(parents=True, exist_ok=True)
    store.save_local(str(out_dir))

    sources_set = {str(m["source"]) for m in all_metadatas}
    meta_payload = {
        "embedding_model": embedding_model,
        "md_files_used": len(sources_set),
        "num_chunks": len(all_chunks),
        "sources": sorted(sources_set),
    }
    with open(out_dir / "disney_faiss_meta.pkl", "wb") as f:
        pickle.dump(meta_payload, f)

    if verbose:
        print(
            f"FAISS 索引已写入: {out_dir}\n"
            f"  嵌入模型: {embedding_model}\n"
            f"  参与 .md 文件数: {meta_payload['md_files_used']}\n"
            f"  向量块数: {len(all_chunks)}\n"
        )

    return FaissBuildResult(
        save_dir=str(out_dir),
        embedding_model=embedding_model,
        md_files_used=int(meta_payload["md_files_used"]),
        num_chunks=len(all_chunks),
        faiss_store=store,
    )


def load_disney_faiss_index(
    faiss_dir: Optional[Union[str, Path]] = None,
    *,
    knowledge_dir: Optional[Union[str, Path]] = None,
    embedding_model: str = "text-embedding-v4",
    dashscope_api_key: Optional[str] = None,
) -> Any:
    """
    Load a FAISS store built by :func:`build_faiss_from_docx_markdown`.

    Args:
        faiss_dir: Directory containing ``index.faiss`` / ``index.pkl``. If omitted,
            uses ``<knowledge_dir>/disney_docx_faiss``.
        knowledge_dir: Used only when ``faiss_dir`` is ``None`` (defaults next to this module).
        embedding_model: Must match the model used when building the index.
        dashscope_api_key: Optional; defaults to ``DASHSCOPE_API_KEY``.
    """
    from langchain_community.embeddings import DashScopeEmbeddings
    from langchain_community.vectorstores import FAISS

    api_key = dashscope_api_key or os.getenv("DASHSCOPE_API_KEY")
    if not api_key:
        raise ValueError("请设置环境变量 DASHSCOPE_API_KEY，或通过 dashscope_api_key 传入。")

    base = Path(__file__).resolve().parent
    kd = Path(knowledge_dir) if knowledge_dir is not None else base / "disney_knowledge_base"
    kd = kd.resolve()

    path = (
        Path(faiss_dir).expanduser().resolve()
        if faiss_dir is not None
        else kd / _DEFAULT_DISNEY_FAISS_DIRNAME
    )

    embeddings = DashScopeEmbeddings(
        model=embedding_model,
        dashscope_api_key=api_key,
    )
    return FAISS.load_local(
        str(path),
        embeddings,
        allow_dangerous_deserialization=True,
    )


def _escape_md_cell(text: str) -> str:
    return text.replace("\\", "\\\\").replace("|", "\\|").replace("\n", "<br>")


def _table_to_markdown(table: Table) -> str:
    rows_data: List[List[str]] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            parts = [p.text for p in cell.paragraphs]
            text = "\n".join(parts).strip()
            cells.append(_escape_md_cell(text) if text else " ")
        rows_data.append(cells)

    if not rows_data:
        return ""

    num_cols = max(len(r) for r in rows_data)
    norm: List[List[str]] = []
    for r in rows_data:
        padded = r + [" "] * (num_cols - len(r))
        norm.append(padded[:num_cols])

    header = norm[0]
    sep = ["---"] * num_cols
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for body_row in norm[1:]:
        lines.append("| " + " | ".join(body_row) + " |")
    return "\n".join(lines) + "\n"


def _split_paragraph_text_by_internal_breaks(
    paragraph: Paragraph, start_page: int
) -> Tuple[List[Tuple[int, str]], int]:
    """
    Split paragraph text into (page, text) segments when Word inserted
    lastRenderedPageBreak or explicit page breaks inside the paragraph.
    """
    page = start_page
    buffers: List[str] = []
    chunks: List[Tuple[int, str]] = []

    def flush():
        nonlocal buffers
        joined = "".join(buffers).strip()
        buffers = []
        if joined:
            chunks.append((page, joined))

    for run in paragraph.runs:
        for el in run._element:
            tag = el.tag
            if tag.endswith("}lastRenderedPageBreak"):
                flush()
                page += 1
            elif tag.endswith("}br") and el.get(qn("w:type")) == "page":
                flush()
                page += 1
            elif tag.endswith("}t"):
                if el.text:
                    buffers.append(el.text)
                if el.tail:
                    buffers.append(el.tail)

    flush()

    if not chunks and paragraph.text and paragraph.text.strip():
        chunks.append((page, paragraph.text.strip()))

    ending_page = page
    return chunks, ending_page


def _estimate_page_from_char_index(char_index: int, chars_per_page: int) -> int:
    if chars_per_page <= 0:
        return 1
    return max(1, char_index // chars_per_page + 1)


def _is_word_lock_or_temp_docx(path: Path) -> bool:
    """
    Word / Office may create lock or temp siblings next to real ``.docx`` files.
    They match ``*.docx`` but are not valid OOXML packages (open raises PackageNotFoundError).
    """
    name = path.name
    return name.startswith("~$") or name.startswith(".~")


def _document_has_page_marks(docx_path: Path) -> bool:
    """Detect pagination hints stored in ``word/document.xml``."""
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return (
        "lastRenderedPageBreak" in xml
        or 'w:type="page"' in xml
        or "w:type='page'" in xml
    )


def iter_document_blocks(document: DocumentObject) -> Iterator[Union[Paragraph, Table]]:
    """Yield paragraphs and top-level tables in document order (skips sectPr)."""
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _table_to_plain_text(table: Table) -> str:
    """Table as plain lines (tabs between cells), aligned with DOCX cell text (not Markdown)."""
    lines: List[str] = []
    for row in table.rows:
        cells: List[str] = []
        for cell in row.cells:
            parts = [p.text for p in cell.paragraphs]
            text = "\n".join(parts).strip().replace("\t", " ")
            cells.append(" ".join(text.splitlines()) if text else "")
        lines.append("\t".join(cells))
    return "\n".join(lines).strip()


def _iter_docx_export_pieces(
    docx_path: Path,
    *,
    estimate_pages_by_chars: bool,
    chars_per_page: int,
) -> Iterator[Tuple[int, int, Literal["break", "estimated", "default"], ChunkType, str, str]]:
    """
    Walk one DOCX with the same chunk_index / page rules as :func:`export_docx_folder_to_markdown`.

    Yields ``(chunk_index, eff_page, page_source, chunk_type, md_body, plain_body)``.
    For paragraphs ``md_body == plain_body``; for tables ``md_body`` is Markdown,
    ``plain_body`` is :func:`_table_to_plain_text`.
    """
    doc = Document(str(docx_path))
    chunk_index = 0
    current_page = 1
    had_any_break_hint = _document_has_page_marks(docx_path)
    linear_char_cursor = 0

    for block in iter_document_blocks(doc):
        if isinstance(block, Paragraph):
            sub_chunks, ending_page = _split_paragraph_text_by_internal_breaks(
                block, current_page
            )
            current_page = ending_page

            for seg_page, text in sub_chunks:
                if not text.strip():
                    continue
                chunk_index += 1
                page_src: Literal["break", "estimated", "default"]
                if estimate_pages_by_chars and not had_any_break_hint:
                    eff_page = _estimate_page_from_char_index(
                        linear_char_cursor, chars_per_page
                    )
                    page_src = "estimated"
                elif had_any_break_hint:
                    eff_page = seg_page
                    page_src = "break"
                else:
                    page_src = "default"
                    eff_page = 1

                linear_char_cursor += len(text)
                yield (chunk_index, eff_page, page_src, "text", text, text)

        elif isinstance(block, Table):
            md_table = _table_to_markdown(block)
            if not md_table.strip():
                continue
            chunk_index += 1
            if estimate_pages_by_chars and not had_any_break_hint:
                eff_page = _estimate_page_from_char_index(
                    linear_char_cursor, chars_per_page
                )
                page_src = "estimated"
            elif had_any_break_hint:
                eff_page = current_page
                page_src = "break"
            else:
                eff_page = 1
                page_src = "default"

            linear_char_cursor += len(md_table)
            plain_table = _table_to_plain_text(block)
            yield (chunk_index, eff_page, page_src, "table", md_table, plain_table)


def _build_docx_plain_chunk_maps(
    docx_path: Path,
    *,
    estimate_pages_by_chars: bool = False,
    chars_per_page: int = 1800,
) -> Tuple[Dict[int, str], Dict[int, int]]:
    """``chunk_index`` → plain text / page, for retrieval materialization from DOCX."""
    plain: Dict[int, str] = {}
    pages: Dict[int, int] = {}
    for idx, pg, _src, _ctype, _md, body_plain in _iter_docx_export_pieces(
        docx_path,
        estimate_pages_by_chars=estimate_pages_by_chars,
        chars_per_page=chars_per_page,
    ):
        plain[idx] = body_plain
        pages[idx] = pg
    return plain, pages


_CHUNK_INDEX_FROM_COMMENT_RE = re.compile(r"chunk_index=(\d+)")


def _chunk_indices_from_md_export_segment(text: str) -> List[int]:
    """Parse ``chunk_index=`` from export HTML comments (may be multiple in one FAISS piece)."""
    return [int(m.group(1)) for m in _CHUNK_INDEX_FROM_COMMENT_RE.finditer(text or "")]


def _materialize_docx_context_for_retrieval(
    docx_path: Path,
    md_segment: str,
    meta: Dict[str, Any],
    *,
    estimate_pages_by_chars: bool = False,
    chars_per_page: int = 1800,
) -> str:
    """
    Replace MD-derived FAISS ``page_content`` with the same logical chunks read from DOCX
    (paragraph text + tables as plain text). Falls back to ``md_segment`` if DOCX is missing
    or cannot be aligned.

    ``estimate_pages_by_chars`` / ``chars_per_page`` must match the values used when
    exporting MD from DOCX if you rely on page-based fallback.
    """
    if not docx_path.is_file():
        return md_segment
    try:
        plain_by_idx, page_by_idx = _build_docx_plain_chunk_maps(
            docx_path,
            estimate_pages_by_chars=estimate_pages_by_chars,
            chars_per_page=chars_per_page,
        )
    except Exception:
        return md_segment

    idxs = sorted(set(_chunk_indices_from_md_export_segment(md_segment)))
    if idxs:
        parts = [plain_by_idx[i] for i in idxs if i in plain_by_idx]
        if parts:
            return "\n\n".join(parts)

    lo, hi = meta.get("page_min"), meta.get("page_max")
    if lo is not None and hi is not None:
        lo_i, hi_i = int(lo), int(hi)
        parts = [
            plain_by_idx[i]
            for i in sorted(plain_by_idx.keys())
            if i in page_by_idx and lo_i <= page_by_idx[i] <= hi_i
        ]
        if parts:
            return "\n\n".join(parts)

    return md_segment


def export_docx_folder_to_markdown(
    knowledge_dir: Optional[Union[str, Path]] = None,
    *,
    estimate_pages_by_chars: bool = False,
    chars_per_page: int = 1800,
    write_md_files: bool = True,
    verbose: bool = True,
) -> List[DocxExportResult]:
    """
    Read every ``*.docx`` under ``disney_knowledge_base``, write sibling ``*.md`` files,
    and return structured chunk metadata (filename, paths, chunk type, page).

    Page numbers:
    - ``break``: from ``w:lastRenderedPageBreak`` / ``w:br w:type="page`` when present.
    - ``estimated``: optional heuristic using ``chars_per_page`` over linear text length.
    - ``default``: ``1`` when no breaks and estimation is off (OOXML often has no layout pages).

    Args:
        knowledge_dir: Folder containing DOCX files. Defaults to ``./disney_knowledge_base``
            next to this module.
        estimate_pages_by_chars: If True and no page-break hints exist, assign estimated pages.
        chars_per_page: Characters per estimated page when estimation is enabled.
        write_md_files: When False, only compute metadata without writing Markdown.
        verbose: Print a short report per file.

    Returns:
        List of :class:`DocxExportResult` with chunk lists and output paths.
    """
    base = Path(__file__).resolve().parent
    folder = Path(knowledge_dir) if knowledge_dir is not None else base / "disney_knowledge_base"
    folder = folder.resolve()

    if not folder.is_dir():
        raise FileNotFoundError(f"Knowledge folder not found: {folder}")

    results: List[DocxExportResult] = []

    for docx_path in sorted(folder.glob("*.docx")):
        if _is_word_lock_or_temp_docx(docx_path):
            if verbose:
                print(f"跳过（Word 临时/锁文件）: {docx_path.name}\n")
            continue
        md_path = docx_path.with_suffix(".md")

        chunks_meta: List[ChunkRecord] = []
        md_parts: List[str] = []
        had_any_break_hint = _document_has_page_marks(docx_path)

        title = docx_path.stem
        md_parts.append(f"# {title}\n\n")

        for chunk_index, eff_page, page_src, ctype, md_body, _plain in _iter_docx_export_pieces(
            docx_path,
            estimate_pages_by_chars=estimate_pages_by_chars,
            chars_per_page=chars_per_page,
        ):
            if ctype == "text":
                preview = md_body if len(md_body) <= 120 else md_body[:117] + "..."
                chunks_meta.append(
                    ChunkRecord(
                        chunk_index=chunk_index,
                        chunk_type="text",
                        page=eff_page,
                        page_source=page_src,
                        content_preview=preview,
                    )
                )
                md_parts.append(
                    f"<!-- chunk_index={chunk_index} type=text page={eff_page} "
                    f"page_source={page_src} -->\n\n{md_body}\n\n"
                )
            else:
                preview = md_body.strip().splitlines()[0][:120]
                chunks_meta.append(
                    ChunkRecord(
                        chunk_index=chunk_index,
                        chunk_type="table",
                        page=eff_page,
                        page_source=page_src,
                        content_preview=preview,
                    )
                )
                md_parts.append(
                    f"<!-- chunk_index={chunk_index} type=table page={eff_page} "
                    f"page_source={page_src} -->\n\n{md_body}\n"
                )

        page_note = ""
        if not had_any_break_hint and not estimate_pages_by_chars:
            page_note = (
                "本文档未包含 OOXML 分页标记（lastRenderedPageBreak / 分页符）；"
                "页码均记为 1。可在 Word 中打开并保存以写入布局分页标记，"
                "或设置 estimate_pages_by_chars=True 启用按字数估算页码。"
            )
        elif estimate_pages_by_chars and not had_any_break_hint:
            page_note = f"页码按约每 {chars_per_page} 字符一页估算。"

        result = DocxExportResult(
            docx_filename=docx_path.name,
            docx_path=str(docx_path),
            md_filename=md_path.name,
            md_path=str(md_path),
            chunks=chunks_meta,
            page_note=page_note,
        )

        if page_note and write_md_files:
            md_parts.append(f"\n---\n\n*导出说明（页码）*: {page_note}\n")

        if write_md_files:
            md_path.write_text("".join(md_parts), encoding="utf-8")

        results.append(result)

        if verbose:
            print_export_summary(result)

    return results


def print_export_summary(result: DocxExportResult) -> None:
    """Print filename, paths, chunk types, and page numbers."""
    print(f"DOCX 文件: {result.docx_filename}")
    print(f"  位置: {result.docx_path}")
    print(f"  MD 输出: {result.md_path}")
    if result.page_note:
        print(f"  页码说明: {result.page_note}")
    for c in result.chunks:
        print(
            f"  chunk {c.chunk_index} | {c.chunk_type:5} | page={c.page} "
            f"({c.page_source}) | {c.content_preview}"
        )
    print()


def export_summary_to_dict(results: List[DocxExportResult]) -> List[Dict[str, Any]]:
    """Serialize export results to plain dicts (e.g. for JSON logging)."""
    out: List[Dict[str, Any]] = []
    for r in results:
        out.append(
            {
                "docx_filename": r.docx_filename,
                "docx_path": r.docx_path,
                "md_filename": r.md_filename,
                "md_path": r.md_path,
                "page_note": r.page_note,
                "chunks": [
                    {
                        "chunk_index": c.chunk_index,
                        "chunk_type": c.chunk_type,
                        "page": c.page,
                        "page_source": c.page_source,
                        "content_preview": c.content_preview,
                    }
                    for c in r.chunks
                ],
            }
        )
    return out


_CLIP_FAISS_INDEX_NAME = "clip_index.faiss"
_CLIP_META_NAME = "clip_metadata.pkl"


def build_clip_image_text_faiss_index(
    ocr_results: List[ImageOcrResult],
    *,
    save_dir: Union[str, Path],
    clip_model_name: str = "openai/clip-vit-base-patch32",
    clip_device: Optional[str] = None,
    verbose: bool = True,
) -> ClipFaissBuildResult:
    """
    Embed each knowledge image with CLIP ``get_image_features``, embed OCR transcript (Chinese + English)
    with CLIP ``get_text_features`` when non-empty; L2-normalize for cosine retrieval with ``IndexFlatIP``.

    Persists ``clip_index.faiss`` and ``clip_metadata.pkl`` under ``save_dir``.

    Requires: ``pip install transformers torch pillow faiss-cpu`` (+ ``tesseract`` OCR already run to fill
    ``ocr_results``).

    Metadata entries (aligned with row order): ``modality`` (``clip_image`` | ``clip_ocr_text``),
    ``image_filename``, ``image_path``, ``ocr_text``.
    """
    try:
        import numpy as np
        import faiss  # type: ignore[import-untyped]
        import torch
        from transformers import CLIPModel, CLIPProcessor
    except ImportError as e:
        raise RuntimeError(
            "CLIP + FAISS 需要: pip install transformers torch faiss-cpu numpy"
        ) from e

    if Image is None:
        raise RuntimeError("需要 Pillow。")

    out = Path(save_dir).expanduser().resolve()
    out.mkdir(parents=True, exist_ok=True)

    os.environ.setdefault("OMP_NUM_THREADS", "1")
    torch.set_num_threads(1)

    device = clip_device or ("cuda" if torch.cuda.is_available() else "cpu")
    proc = CLIPProcessor.from_pretrained(clip_model_name)
    model = CLIPModel.from_pretrained(clip_model_name).to(device)
    model.eval()

    vectors: List[Any] = []
    records: List[Dict[str, Any]] = []
    n_img = 0
    n_txt = 0

    def _norm_np(x: Any) -> Any:
        import numpy as np

        x = np.asarray(x, dtype=np.float32).reshape(-1)
        n = float(np.linalg.norm(x))
        if n < 1e-12:
            return x
        return (x / n).astype(np.float32)

    with torch.no_grad():
        for rec in ocr_results:
            path = Path(rec.image_path)
            if not path.is_file():
                continue
            pil = Image.open(path).convert("RGB")
            ins = proc(images=pil, return_tensors="pt")
            ins = {k: v.to(device) for k, v in ins.items()}
            img_feat = model.get_image_features(**ins)
            img_feat = img_feat / img_feat.norm(dim=-1, keepdim=True)
            v_img = _norm_np(img_feat.cpu().numpy()[0])
            vectors.append(v_img)
            records.append(
                {
                    "modality": "clip_image",
                    "image_filename": rec.image_filename,
                    "image_path": rec.image_path,
                    "ocr_text": rec.ocr_text,
                }
            )
            n_img += 1

            tx = (rec.ocr_text or "").strip()
            if tx and tx != "*(未识别到文本)*":
                tins = proc(
                    text=[tx],
                    return_tensors="pt",
                    padding=True,
                    truncation=True,
                    max_length=77,
                )
                tins = {k: v.to(device) for k, v in tins.items()}
                txt_feat = model.get_text_features(**tins)
                txt_feat = txt_feat / txt_feat.norm(dim=-1, keepdim=True)
                v_txt = _norm_np(txt_feat.cpu().numpy()[0])
                vectors.append(v_txt)
                records.append(
                    {
                        "modality": "clip_ocr_text",
                        "image_filename": rec.image_filename,
                        "image_path": rec.image_path,
                        "ocr_text": tx,
                    }
                )
                n_txt += 1

    if not vectors:
        raise ValueError("没有可用的图片向量（检查路径与 OCR 结果是否为空）。")

    dim = int(vectors[0].shape[0])
    mat = np.stack(vectors, axis=0).astype("float32")
    index = faiss.IndexFlatIP(dim)
    index.add(mat)

    faiss.write_index(index, str(out / _CLIP_FAISS_INDEX_NAME))
    payload = {
        "records": records,
        "clip_model_name": clip_model_name,
        "embedding_dim": dim,
        "metric": "inner_product_on_l2_normalized",
    }
    with open(out / _CLIP_META_NAME, "wb") as f:
        pickle.dump(payload, f)

    if verbose:
        print(
            f"CLIP+FAISS 已写入: {out}\n"
            f"  模型: {clip_model_name}\n"
            f"  向量条数: {len(records)}（图像 {n_img} + OCR 文本 {n_txt}），维度 {dim}\n"
        )

    return ClipFaissBuildResult(
        save_dir=str(out),
        clip_model_name=clip_model_name,
        embedding_dim=dim,
        num_vectors=len(records),
        num_image_vectors=n_img,
        num_text_vectors=n_txt,
    )


def load_clip_image_text_faiss(
    faiss_dir: Optional[Union[str, Path]] = None,
    *,
    knowledge_dir: Optional[Union[str, Path]] = None,
) -> Tuple[Any, List[Dict[str, Any]], Dict[str, Any]]:
    """
    Load CLIP+FAISS bundle written by :func:`build_clip_image_text_faiss_index`.

    Returns:
        ``(faiss_index, records, payload_extras)`` where ``records`` is the per-vector metadata list.
    """
    try:
        import faiss  # type: ignore[import-untyped]
    except ImportError as e:
        raise RuntimeError("需要 pip install faiss-cpu") from e

    base = Path(__file__).resolve().parent
    kd = Path(knowledge_dir) if knowledge_dir is not None else base / "disney_knowledge_base"
    kd = kd.resolve()
    path = (
        Path(faiss_dir).expanduser().resolve()
        if faiss_dir is not None
        else kd / "clip_image_text_faiss"
    )
    idx_path = path / _CLIP_FAISS_INDEX_NAME
    meta_path = path / _CLIP_META_NAME
    if not idx_path.is_file() or not meta_path.is_file():
        raise FileNotFoundError(f"未找到 CLIP FAISS 文件: {idx_path} / {meta_path}")

    index = faiss.read_index(str(idx_path))
    with open(meta_path, "rb") as f:
        payload = pickle.load(f)
    records = payload.get("records", [])
    return index, records, payload


_IMAGE_ROUTE_PATTERN = re.compile(
    r"图片|图像|海报|照片|截图|画面上|图中|看图|海报上|视觉上|"
    r"什么样子|长什么样|哪张图|这张图|万圣节.*图|插画|封面图",
)

_QWEN_DEFAULT = "qwen-turbo"


def _strip_paired_markdown_bold(text: str) -> str:
    """Remove ``**...**`` from model output (plain-text terminals do not render Markdown)."""
    if not text:
        return text
    out = text
    prev = None
    while prev != out:
        prev = out
        out = re.sub(r"\*\*([^*]+)\*\*", r"\1", out)
    return out


def _page_numbers_from_md_export_chunk(text: str) -> Tuple[List[int], Optional[int]]:
    """
    Collect ``page=N`` integers from DOCX-export HTML comments (``<!-- ... page=N ... -->``).
    Returns ``(all_page_numbers, dominant_page)``.
    """
    nums = [int(m.group(1)) for m in re.finditer(r"\bpage=(\d+)", text or "")]
    if not nums:
        return [], None
    from collections import Counter

    dominant = Counter(nums).most_common(1)[0][0]
    return nums, dominant


def _page_label_from_numbers(nums: List[int]) -> str:
    if not nums:
        return ""
    lo, hi = min(nums), max(nums)
    return f"第{lo}页" if lo == hi else f"第{lo}–{hi}页"


def _resolve_doc_citation_page_label(meta: Dict[str, Any], content: str) -> str:
    """Human-readable page label for citations (file + page display)."""
    nums, _ = _page_numbers_from_md_export_chunk(content)
    lbl = _page_label_from_numbers(nums)
    if lbl:
        return lbl
    pl = meta.get("page_label")
    if isinstance(pl, str) and pl.strip():
        return pl.strip()
    if meta.get("page") is not None:
        pmn = meta.get("page_min")
        pmx = meta.get("page_max")
        if pmn is not None and pmx is not None and pmn != pmx:
            return f"第{pmn}–{pmx}页"
        return f"第{int(meta['page'])}页"
    return "页码未标注"


def route_query_to_kb_kind(
    query: str,
    *,
    force: Optional[Literal["text_knowledge", "image_knowledge"]] = None,
) -> Literal["text_knowledge", "image_knowledge"]:
    """
    Heuristic route: mentions of visuals → CLIP+FAISS (image_knowledge); else DOCX‑MD embeddings.
    Override with ``force`` when routing is explicit.
    """
    if force is not None:
        return force
    q = (query or "").strip()
    if _IMAGE_ROUTE_PATTERN.search(q):
        return "image_knowledge"
    return "text_knowledge"


def _clip_encode_query_text(
    question: str,
    clip_model_name: str,
    *,
    clip_device: Optional[str] = None,
) -> Any:
    """Return L2-normalized CLIP text embedding (numpy float32 vector)."""
    import numpy as np
    import torch
    from transformers import CLIPModel, CLIPProcessor

    os.environ.setdefault("OMP_NUM_THREADS", "1")
    torch.set_num_threads(1)

    device = clip_device or ("cuda" if torch.cuda.is_available() else "cpu")
    proc = CLIPProcessor.from_pretrained(clip_model_name)
    model = CLIPModel.from_pretrained(clip_model_name).to(device)
    model.eval()

    tins = proc(
        text=[question.strip()],
        return_tensors="pt",
        padding=True,
        truncation=True,
        max_length=77,
    )
    tins = {k: v.to(device) for k, v in tins.items()}
    with torch.no_grad():
        feat = model.get_text_features(**tins)
        feat = feat / feat.norm(dim=-1, keepdim=True)
    v = feat.cpu().numpy().astype("float32")[0]
    nv = float(np.linalg.norm(v))
    if nv > 1e-12:
        v = (v / nv).astype(np.float32)
    return v


def _record_to_clip_context(rank: int, rec: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
    """Turn one CLIP metadata row into prompt text + citation dict."""
    modality = rec.get("modality", "")
    fn = rec.get("image_filename", "")
    path = rec.get("image_path", "")
    ocr = (rec.get("ocr_text") or "").strip()

    img_page_note = "（图片素材，无文档页码）"

    if modality == "clip_image":
        label = "图片条目（CLIP 图像向量）"
        excerpt = (ocr[:600] + "…") if len(ocr) > 600 else ocr
        body = (
            f"[{rank}] [{label}] 文件: {fn}\n路径: {path}\n"
            f"页码参照: {img_page_note}\n图像关联 OCR 摘录:\n{excerpt or '(无 OCR)'}"
        )
        cite_line = f"[{rank}] {label}｜{fn}｜{img_page_note}"
    else:
        label = "图中文字条目（CLIP 文本向量 / OCR）"
        body = (
            f"[{rank}] [{label}] 来源图片: {fn}\n路径: {path}\n"
            f"页码参照: {img_page_note}\n文本内容:\n"
            f"{(ocr[:900] + '…') if len(ocr) > 900 else ocr}"
        )
        cite_line = f"[{rank}] {label}｜{fn}｜{img_page_note}"

    cite = {
        "rank": rank,
        "modality": modality,
        "image_filename": fn,
        "image_path": path,
        "page_label": img_page_note,
        "citation_line": cite_line,
    }
    return body, cite


def retrieve_top_contexts_clip(
    question: str,
    *,
    knowledge_dir: Optional[Union[str, Path]] = None,
    clip_faiss_dir: Optional[Union[str, Path]] = None,
    top_k: int = 5,
    clip_device: Optional[str] = None,
) -> Tuple[List[str], List[Dict[str, Any]], str]:
    """CLIP 文本检索 ``clip_image_text_faiss``：返回上下文段落、引用信息、所用的 CLIP 模型名。"""
    import numpy as np

    index, records, payload = load_clip_image_text_faiss(
        clip_faiss_dir,
        knowledge_dir=knowledge_dir,
    )
    if not records:
        raise ValueError("CLIP 索引元数据为空")

    clip_name = str(payload.get("clip_model_name", "openai/clip-vit-base-patch32"))
    qv = _clip_encode_query_text(question, clip_name, clip_device=clip_device)
    qv = np.asarray(qv, dtype=np.float32).reshape(1, -1)

    k = min(top_k, index.ntotal)
    scores, ids = index.search(qv, k)

    paragraphs: List[str] = []
    cites: List[Dict[str, Any]] = []
    order = 1
    for idx_i, score in zip(ids[0].tolist(), scores[0].tolist()):
        if idx_i < 0 or idx_i >= len(records):
            continue
        para, cite = _record_to_clip_context(order, records[idx_i])
        cite["score"] = float(score)
        cite["faiss_index_row"] = int(idx_i)
        paragraphs.append(para)
        cites.append(cite)
        order += 1

    return paragraphs, cites, clip_name


def retrieve_top_contexts_docx_md(
    question: str,
    *,
    knowledge_dir: Optional[Union[str, Path]] = None,
    faiss_dir: Optional[Union[str, Path]] = None,
    top_k: int = 5,
    embedding_model: str = "text-embedding-v4",
    dashscope_api_key: Optional[str] = None,
    estimate_pages_by_chars: bool = False,
    chars_per_page: int = 1800,
    use_docx_source_for_context: bool = True,
) -> Tuple[List[str], List[Dict[str, Any]]]:
    """
    DashScope 嵌入 + LangChain FAISS：返回上下文段落与引用信息。

    向量仍由顶层 ``*.md`` 切分嵌入（与 :func:`build_faiss_from_docx_markdown` 一致）；
    默认将每条命中中的「展示正文」替换为对应 **DOCX 原文**（按导出时的 ``chunk_index``，
    或按元数据中的 ``page_min``/``page_max`` 回退），不再把 Markdown 表格等排版当作模型上下文。

    若导出 MD 时使用了 ``estimate_pages_by_chars=True``，检索时须传相同的
    ``estimate_pages_by_chars`` / ``chars_per_page``，页码回退才能与 DOCX 对齐。
    """
    base = Path(__file__).resolve().parent
    kd = Path(knowledge_dir) if knowledge_dir is not None else base / "disney_knowledge_base"
    kd = kd.resolve()

    vectorstore = load_disney_faiss_index(
        faiss_dir,
        knowledge_dir=knowledge_dir,
        embedding_model=embedding_model,
        dashscope_api_key=dashscope_api_key,
    )
    docs = vectorstore.similarity_search(question, k=top_k)

    paragraphs: List[str] = []
    cites: List[Dict[str, Any]] = []
    for i, d in enumerate(docs, start=1):
        meta = dict(d.metadata or {})
        src = meta.get("source", "?")
        md_pc = (d.page_content or "").strip()
        page_label = _resolve_doc_citation_page_label(meta, d.page_content)

        docx_path: Optional[Path] = None
        if use_docx_source_for_context:
            dp = meta.get("docx_path")
            if isinstance(dp, str) and dp.strip():
                cpp = Path(dp)
                if cpp.is_file():
                    docx_path = cpp
            if docx_path is None and meta.get("md_path"):
                cpp = Path(str(meta["md_path"])).with_suffix(".docx")
                if cpp.is_file():
                    docx_path = cpp
            if docx_path is None and isinstance(src, str) and src.lower().endswith(".md"):
                cpp = kd / f"{Path(src).stem}.docx"
                if cpp.is_file():
                    docx_path = cpp

        body_text = md_pc
        if docx_path is not None:
            body_text = _materialize_docx_context_for_retrieval(
                docx_path,
                md_pc,
                meta,
                estimate_pages_by_chars=estimate_pages_by_chars,
                chars_per_page=chars_per_page,
            )

        docx_name = f"{Path(str(src)).stem}.docx" if isinstance(src, str) else "?"
        docx_line = str(docx_path.resolve()) if docx_path is not None else (meta.get("docx_path") or "")
        body = (
            f"[{i}] [文档片段 text-embedding] 来源 Word: {docx_name}（向量由同名 .md 检索对齐）\n"
            f"DOCX 路径: {docx_line}\n页码参照: {page_label}\n"
            f"内容（DOCX 原文）:\n{body_text.strip()}"
        )
        paragraphs.append(body)
        cites.append(
            {
                "rank": i,
                "modality": "docx_md_chunk",
                "source": src,
                "md_path": meta.get("md_path"),
                "docx_path": str(docx_path.resolve()) if docx_path is not None else meta.get("docx_path"),
                "context_from": "docx" if docx_path is not None else "md_only",
                "chunk_in_file": meta.get("chunk_in_file"),
                "page": meta.get("page"),
                "page_label": page_label,
                "citation_line": f"[{i}] 文档｜{docx_name}｜{page_label}",
            }
        )
    return paragraphs, cites


def rag_answer_disney_question(
    question: str,
    *,
    top_k: int = 5,
    llm_model: str = _QWEN_DEFAULT,
    knowledge_dir: Optional[Union[str, Path]] = None,
    docx_faiss_dir: Optional[Union[str, Path]] = None,
    clip_faiss_dir: Optional[Union[str, Path]] = None,
    embedding_model: str = "text-embedding-v4",
    dashscope_api_key: Optional[str] = None,
    clip_device: Optional[str] = None,
    force_intent: Optional[Literal["text_knowledge", "image_knowledge"]] = None,
    verbose: bool = True,
    estimate_pages_by_chars: bool = False,
    chars_per_page: int = 1800,
    use_docx_source_for_context: bool = True,
) -> RagAnswerResult:
    """
    检索用户问题：根据轻量规则划为「文档知识」或「图片/视觉知识」，在对应 FAISS 中取 top‑K，
    将上下文与问题发给 DashScope **qwen-turbo**（可用 ``llm_model`` 改为同系其它模型），返回答案与出处说明。

    - **text_knowledge**：``build_faiss_from_docx_markdown`` 生成的 ``disney_docx_faiss``（DashScope 向量）。
    - **image_knowledge**：``clip_image_text_faiss``（CLIP 向量，需已构建）。

    环境变量 ``DASHSCOPE_API_KEY`` 必填（嵌入 + LLM）。

    Args:
        question: 用户问题。
        top_k: 检索条数（默认 5）。
        llm_model: DashScope 对话模型名，默认 ``qwen-turbo``。
        force_intent: 强制走文档或图片索引；``None`` 时使用 ``route_query_to_kb_kind`` 启发式。
        verbose: 是否打印检索意图与出处摘要。

    Returns:
        :class:`RagAnswerResult`（``top_contexts`` 含每条引用的 citation_line / 路径等）。
    """
    from langchain_community.llms import Tongyi

    api_key = dashscope_api_key or os.getenv("DASHSCOPE_API_KEY")
    if not api_key:
        raise ValueError("请设置环境变量 DASHSCOPE_API_KEY。")

    if not (question or "").strip():
        raise ValueError("问题不能为空。")

    intent = route_query_to_kb_kind(question, force=force_intent)

    contexts: List[str] = []
    cites: List[Dict[str, Any]] = []
    clip_name = ""

    if intent == "text_knowledge":
        contexts, cites = retrieve_top_contexts_docx_md(
            question,
            knowledge_dir=knowledge_dir,
            faiss_dir=docx_faiss_dir,
            top_k=top_k,
            embedding_model=embedding_model,
            dashscope_api_key=api_key,
            estimate_pages_by_chars=estimate_pages_by_chars,
            chars_per_page=chars_per_page,
            use_docx_source_for_context=use_docx_source_for_context,
        )
    else:
        try:
            contexts, cites, clip_name = retrieve_top_contexts_clip(
                question,
                knowledge_dir=knowledge_dir,
                clip_faiss_dir=clip_faiss_dir,
                top_k=top_k,
                clip_device=clip_device,
            )
        except FileNotFoundError:
            if verbose:
                print(
                    "[WARN] 未找到 CLIP FAISS，已回退到文档向量库。",
                    flush=True,
                )
            intent = "text_knowledge"
            contexts, cites = retrieve_top_contexts_docx_md(
                question,
                knowledge_dir=knowledge_dir,
                faiss_dir=docx_faiss_dir,
                top_k=top_k,
                embedding_model=embedding_model,
                dashscope_api_key=api_key,
                estimate_pages_by_chars=estimate_pages_by_chars,
                chars_per_page=chars_per_page,
                use_docx_source_for_context=use_docx_source_for_context,
            )

    if not contexts:
        raise RuntimeError("未检索到任何上下文；请先构建对应 FAISS 索引。")

    joined_ctx = "\n\n---\n\n".join(contexts)

    prompt = (
        "你是迪士尼园区知识助手。下面「参考资料」来自知识库检索，请优先依据资料回答；"
        "若资料不足以回答，请说明无法回答，并避免编造。\n\n"
        f"【参考资料】\n{joined_ctx}\n\n"
        f"【用户问题】\n{question.strip()}\n\n"
        "【回答要求】\n"
        "1. 直接回答用户问题，语气友好，对于比较难以理解的可以进行适当的举例说明。\n"
        "2. 在回答末尾用单独一行写「出处：」，列出用到的编号（如 [1][2]），"
        "每条须同时写明来源文件名与页码（文档类资料已标在「页码参照」；"
        "图片类写「图片、无文档页码」）。\n"
        "3. 全文使用纯中文说明即可，禁止使用 Markdown 或其它标记语法"
        "（尤其不要用星号成对包裹词语来模拟加粗）。\n"
    )

    llm = Tongyi(model_name=llm_model, dashscope_api_key=api_key)
    inv = getattr(llm, "invoke", None)
    answer = inv(prompt) if callable(inv) else llm(prompt)
    answer = _strip_paired_markdown_bold(answer if isinstance(answer, str) else str(answer))

    summary_lines = []
    for c in cites:
        summary_lines.append(c.get("citation_line", ""))
    src_summary = "\n".join(summary_lines)

    if verbose:
        route_label = "文档向量（text-embedding）" if intent == "text_knowledge" else "CLIP 向量"
        print(f"[RAG] 检索通道: {route_label} ({intent})", flush=True)
        if clip_name:
            print(f"[RAG] CLIP 模型: {clip_name}", flush=True)
        print(f"[RAG] 引用条目数: {len(cites)}", flush=True)

    return RagAnswerResult(
        answer=answer,
        intent=intent,
        llm_model=llm_model,
        top_contexts=cites,
        prompt_sources_summary=src_summary,
    )


def export_images_folder_to_markdown_ocr(
    images_dir: Optional[Union[str, Path]] = None,
    *,
    lang: str = "chi_sim+eng",
    recursive: bool = False,
    write_md_files: bool = True,
    verbose: bool = True,
    tesseract_cmd: Optional[str] = None,
    preprocess: bool = True,
    min_long_edge: int = 1800,
    max_long_edge: Optional[int] = 4500,
    contrast: float = 1.92,
    sharpness: float = 1.78,
    unsharp_mask: bool = True,
    apply_median: bool = False,
    binarize: bool = False,
    binarize_threshold: int = 145,
    oem: int = 3,
    psm: int = 3,
    preserve_interword_spaces: bool = True,
    tesseract_extra_config: str = "",
    clip_faiss_save_dir: Optional[Union[str, Path, bool]] = None,
    clip_model_name: str = "openai/clip-vit-base-patch32",
    clip_device: Optional[str] = None,
) -> List[ImageOcrResult]:
    """
    OCR all images under ``disney_knowledge_base/images`` (by default) with pytesseract,
    and write sibling ``.md`` files (same basename as each image).

    When ``clip_faiss_save_dir`` is set (``True`` / ``\"auto\"`` / a path), after OCR completes,
    builds a **CLIP** image + OCR-text **FAISS** index under ``disney_knowledge_base/clip_image_text_faiss``
    (for ``auto`` / ``True``) or under the given directory. Same embedding space as
    ``openai/clip-vit-base-patch32``: one vector per image plus one per non-empty OCR text block.

    Defaults are tuned for **simplified Chinese + English** in one pass (``chi_sim+eng``):
    larger upscaling for thin Latin glyphs, slightly stronger contrast/sharpness, optional
    UnsharpMask, ``psm=3`` for typical posters (multiple text blocks), and
    ``preserve_interword_spaces`` so English words keep spaces.

    Requires system ``tesseract`` and Python packages ``pytesseract``, ``Pillow``.
    Install language data: ``chi_sim`` and ``eng`` (e.g. macOS:
    ``brew install tesseract-lang``). Use ``chi_tra+eng`` for Traditional Chinese.

    **Recognition tips** (when results are still poor):

    - Dense single column only: try ``psm=6``.
    - Sparse / scattered labels: try ``psm=11``.
    - Very light text: raise ``contrast`` or ``binarize=True`` (may hurt smooth fonts).
    - Noise: ``apply_median=True`` (may blur thin strokes).

    Output MD includes a short metadata block: image name, absolute path, MIME / detected type,
    then the recognized text.

    Args:
        images_dir: Folder containing images. Defaults to ``disney_knowledge_base/images``
            next to this module.
        lang: Tesseract ``-l`` languages. Use ``chi_sim+eng`` for 简中+英文；``chi_tra+eng`` for 繁中+英文.
        recursive: If True, walk subfolders for supported image extensions.
        write_md_files: When False, run OCR but do not write Markdown.
        verbose: Print one line per image (name, path, type).
        tesseract_cmd: Path to ``tesseract`` binary; default tries PATH then Homebrew paths.
        preprocess: If True, grayscale / upscale / autocontrast / sharpen before OCR.
        min_long_edge: If longest side is below this (px), scale image up before OCR (higher helps EN+CJK).
        max_long_edge: If set, downscale very large images (speed / memory cap).
        contrast / sharpness: Pillow enhancement factors after autocontrast.
        unsharp_mask: Light UnsharpMask after sharpen (helps Latin edges when image was upscaled).
        apply_median: Apply 3×3 median denoise (optional).
        binarize: Simple black/white threshold after enhancements (optional).
        binarize_threshold: Pixel threshold for ``binarize`` (0–255).
        oem: Tesseract OCR Engine Mode (default 3 = LSTM).
        psm: Page Segmentation Mode. Default 3 = fully automatic (good for mixed blocks); 6 = single block.
        preserve_interword_spaces: Pass ``-c preserve_interword_spaces=1`` for clearer English spacing.
        tesseract_extra_config: Extra CLI flags appended to config.
        clip_faiss_save_dir: If truthy, build CLIP+FAISS after OCR. Use ``True`` / ``\"auto\"`` for
            ``<parent of images>/clip_image_text_faiss`` (parent is typically ``disney_knowledge_base``).
        clip_model_name: Hugging Face CLIP checkpoint for image/text encoders.
        clip_device: Torch device string; default picks CUDA when available.

    Returns:
        List of :class:`ImageOcrResult`.
    """
    if Image is None or pytesseract is None:
        raise RuntimeError(
            "需要安装 Pillow 与 pytesseract：pip install pytesseract pillow；"
            "并在系统安装 Tesseract OCR（macOS: brew install tesseract tesseract-lang）。"
        )

    _resolve_tesseract_cmd(tesseract_cmd)

    base = Path(__file__).resolve().parent
    folder = Path(images_dir) if images_dir is not None else base / "disney_knowledge_base" / "images"
    folder = folder.resolve()

    if not folder.is_dir():
        raise FileNotFoundError(f"图片目录不存在: {folder}")

    def _collect_files() -> List[Path]:
        if recursive:
            paths: List[Path] = []
            for p in folder.rglob("*"):
                if p.is_file() and p.suffix.lower() in _IMAGE_SUFFIXES:
                    paths.append(p)
            return sorted(paths)
        return sorted(
            p
            for p in folder.iterdir()
            if p.is_file() and p.suffix.lower() in _IMAGE_SUFFIXES
        )

    image_paths = _collect_files()
    results: List[ImageOcrResult] = []

    for image_path in image_paths:
        mime_guess, _ = mimetypes.guess_type(str(image_path))
        mime = mime_guess or "application/octet-stream"

        with Image.open(image_path) as im_copy:
            im = im_copy.copy()
            pil_format = (im.format or image_path.suffix.upper().lstrip(".")) or "UNKNOWN"
            work = (
                _preprocess_image_for_ocr(
                    im,
                    min_long_edge=min_long_edge,
                    max_long_edge=max_long_edge,
                    contrast=contrast,
                    sharpness=sharpness,
                    apply_median=apply_median,
                    binarize=binarize,
                    threshold=binarize_threshold,
                    unsharp_mask=unsharp_mask,
                )
                if preprocess
                else im
            )
            tess_cfg = _build_tesseract_config(
                oem=oem,
                psm=psm,
                extra=tesseract_extra_config,
                config_vars=_default_zh_en_tesseract_vars(
                    preserve_interword_spaces=preserve_interword_spaces
                ),
            )
            text = pytesseract.image_to_string(work, lang=lang, config=tess_cfg)

        text = text.strip()
        md_path = image_path.with_suffix(".md")

        md_lines = [
            f"# {image_path.stem}\n\n",
            "## 图片来源信息\n\n",
            f"- **图片名称**: {image_path.name}\n",
            f"- **位置**: {image_path.resolve()}\n",
            f"- **图片类型**: {mime}（PIL 格式: {pil_format}）\n\n",
            "---\n\n",
            "## OCR 文本\n\n",
            text if text else "*(未识别到文本)*",
            "\n",
        ]
        md_body = "".join(md_lines)

        rec = ImageOcrResult(
            image_filename=image_path.name,
            image_path=str(image_path.resolve()),
            image_type=mime,
            md_filename=md_path.name,
            md_path=str(md_path.resolve()),
            ocr_text=text,
        )
        results.append(rec)

        if write_md_files:
            md_path.write_text(md_body, encoding="utf-8")

        if verbose:
            print(f"图片: {rec.image_filename}")
            print(f"  位置: {rec.image_path}")
            print(f"  图片类型: {rec.image_type}（{pil_format}）")
            print(f"  MD: {rec.md_path}\n")

    _clip_dir: Optional[Path] = None
    if clip_faiss_save_dir is not None and clip_faiss_save_dir is not False:
        if clip_faiss_save_dir is True or clip_faiss_save_dir == "auto":
            _clip_dir = folder.parent / "clip_image_text_faiss"
        else:
            _clip_dir = Path(str(clip_faiss_save_dir)).expanduser().resolve()
        if results:
            build_clip_image_text_faiss_index(
                results,
                save_dir=_clip_dir,
                clip_model_name=clip_model_name,
                clip_device=clip_device,
                verbose=verbose,
            )

    return results


def image_ocr_results_to_dict(results: List[ImageOcrResult]) -> List[Dict[str, Any]]:
    """Serialize OCR results to plain dicts."""
    return [
        {
            "image_filename": r.image_filename,
            "image_path": r.image_path,
            "image_type": r.image_type,
            "md_filename": r.md_filename,
            "md_path": r.md_path,
            "ocr_text": r.ocr_text,
        }
        for r in results
    ]


if __name__ == "__main__":
    OnlyQuery = True

    if not OnlyQuery:
        export_docx_folder_to_markdown()
        # 先构建 DOCX→MD 的 DashScope FAISS：避免后面 CLIP 下载失败 / segfault 时整段脚本未执行到此处
        build_faiss_from_docx_markdown()
        try:
            export_images_folder_to_markdown_ocr(clip_faiss_save_dir="auto")
        except Exception as exc:
            print(f"[WARN] 图片 OCR 或 CLIP+FAISS 未完成（不影响已生成的 disney_docx_faiss）：{exc}")

    out = rag_answer_disney_question("迪士尼老人票年龄要求是多少？")
    print(out.answer)
    print(out.prompt_sources_summary)

    out2 = rag_answer_disney_question(
    "万圣节海报上写了什么？",
    force_intent="image_knowledge",
    )
    print(out2.answer)
    print(out2.prompt_sources_summary)

    out3 = rag_answer_disney_question("我想了解一下迪士尼门票的退款流程")
    print(out3.answer)
    print(out3.prompt_sources_summary)